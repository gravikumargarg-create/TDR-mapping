"""
Generate RUNBOOK.docx from the runbook content for sharing with the team.
Run: pip install python-docx && python generate_runbook.py
Output: RUNBOOK.docx in the same folder.
"""
import os

try:
    from docx import Document
except ImportError:
    print("Install python-docx: pip install python-docx")
    raise

def add_para(doc, text, bold=False):
    p = doc.add_paragraph(text)
    if bold:
        for run in p.runs:
            run.bold = True
    return p

def main():
    doc = Document()
    doc.add_heading("TDR Mapping Sheet Creation – Runbook", level=0)
    add_para(doc, "Version: 1.0", bold=True)
    add_para(doc, "Tool: TDR Streamlit App (TDR mapping sheet creation)")
    add_para(doc, "Audience: Team members who run the tool or need to support it.")
    doc.add_paragraph()

    doc.add_heading("1. Purpose", level=1)
    doc.add_paragraph(
        "The TDR Mapping Sheet Creation tool:"
    )
    doc.add_paragraph(
        "• Reads TDR Data (Excel with TDR-###### sections and BANs), LVT Report (Excel with BAN-wise status, e.g. BAN Wise Result), and optionally Device Details (Excel with CUSTOMER_ID and device columns)."
    )
    doc.add_paragraph(
        "• Produces: (1) A main report Excel with TDR Info (TDR, BAN, Status, Failure Description, Check ID, Comments) and TDR Summary; "
        "(2) Per-TDR Excel files (one per TDR) in a ZIP, each with the TDR section and optionally a Device Details sheet when Device Details Excel is provided."
    )
    doc.add_paragraph(
        "• Optionally adds Device Details to each TDR-wise file by matching BAN/Customer ID from an uploaded Device Details Excel."
    )
    doc.add_paragraph()

    doc.add_heading("2. How to Access", level=1)
    doc.add_paragraph("• Streamlit Cloud (recommended): Use the shared app URL (e.g. from your team or deployment).")
    doc.add_paragraph("• Local run:")
    doc.add_paragraph("  cd TDR-mapping", style="List Bullet")
    doc.add_paragraph("  pip install -r requirements.txt", style="List Bullet")
    doc.add_paragraph("  streamlit run app.py", style="List Bullet")
    doc.add_paragraph("  Then open the URL shown in the terminal (e.g. http://localhost:8501).")
    doc.add_paragraph()

    doc.add_heading("3. Prerequisites", level=1)
    doc.add_paragraph("• TDR Data – Excel (.xlsx/.xlsm) with at least one sheet containing TDR section headers (e.g. TDR-202958) and 9-digit BAN IDs.")
    doc.add_paragraph("• LVT Report – Excel with a sheet named BAN Wise Result (or similar; spaces/underscores allowed) containing BAN and Status (e.g. Passed/Failed).")
    doc.add_paragraph("• Device Details (optional) – Excel with a sheet whose first row has CUSTOMER_ID (or BAN/Customer ID) and device columns (MSISDN, IMEI, ESN, EID, DEVICE_MODEL, UICCID, UIMSI, TIMSI, DEVICE_LOCK_STATUS).")
    doc.add_paragraph()

    doc.add_heading("4. Input Options", level=1)
    doc.add_heading("4.1 TDR file from", level=2)
    doc.add_paragraph("• Local file: TDR Data will be chosen from the files you upload.")
    doc.add_paragraph("• SharePoint: If configured, you can pick the TDR file directly from the SharePoint folder. Otherwise, use the link to open the folder, download the file, and upload it with the other files.")
    doc.add_heading("4.2 Upload file(s)", level=2)
    doc.add_paragraph("• One file upload accepts multiple files. Upload one or more Excel files in a single step.")
    doc.add_paragraph("• The app automatically detects: TDR Data (first file with a sheet with TDR sections), LVT Report (first file with a sheet named BAN Wise Result), Device Details (first file with a sheet with CUSTOMER_ID-style column).")
    doc.add_paragraph("• After detection you get dropdowns to choose which sheet to use for TDR Data and LVT Report; Device Details uses the first sheet with CUSTOMER_ID.")
    doc.add_paragraph()

    doc.add_heading("5. Steps to Run", level=1)
    steps = [
        "Choose TDR file from: Local file or SharePoint (if available).",
        "Upload one or more Excel files (TDR Data, LVT Report, and/or Device Details).",
        "Confirm the detected TDR Data and LVT Report files and select the sheets from the dropdowns.",
        "If Device Details was detected, it will be used for the Device Details sheet in each TDR-wise file.",
        "Click Run TDR.",
        "When processing finishes: Download main report and Download per-TDR (ZIP).",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}", style="List Number")
    doc.add_paragraph()

    doc.add_heading("6. Outputs", level=1)
    doc.add_heading("6.1 Main report Excel", level=2)
    doc.add_paragraph("TDR Info sheet: Columns TDR, BAN, Status, Failure Description, Check ID, Comments. Status from LVT; Failure Description and Check ID from BAN Wise Failures sheet (if present); for Passed rows these show N/A. TDR Summary sheet: per-TDR counts and overall summary.")
    doc.add_heading("6.2 Per-TDR Excel files (in ZIP)", level=2)
    doc.add_paragraph("One file per TDR. First sheet: TDR section data. Device Details sheet (if Device Details Excel was uploaded): rows where CUSTOMER_ID matches that TDR's BANs, sorted by CUSTOMER_ID ascending; long numeric columns stored as text.")
    doc.add_paragraph()

    doc.add_heading("7. Device Details", level=1)
    doc.add_paragraph("Source: Optional Excel with CUSTOMER_ID (or equivalent) and device columns. For each TDR, the tool copies matching rows into the Device Details sheet of that TDR's Excel, sorted by CUSTOMER_ID ascending. Long IDs are written as text to avoid truncation.")
    doc.add_paragraph()

    doc.add_heading("8. SharePoint (optional)", level=1)
    doc.add_paragraph("If Sites.Read.All is granted for the app, you can choose TDR file from SharePoint and pick the TDR file from a dropdown. Secrets (AZURE_TENANT_ID, AZURE_CLIENT_ID, AZURE_CLIENT_SECRET) must be set in app settings; do not commit secrets to code or Git.")
    doc.add_paragraph()

    doc.add_heading("9. Troubleshooting", level=1)
    doc.add_paragraph("No TDR Data detected → Ensure at least one uploaded file has a sheet with TDR-###### headers and 9-digit BANs.")
    doc.add_paragraph("No LVT Report detected → Ensure at least one file has a sheet named BAN Wise Result.")
    doc.add_paragraph("Status always Not found → Check LVT sheet has BAN and Status columns and BAN values match TDR BANs.")
    doc.add_paragraph("Device Details sheet empty → Confirm Device Details Excel has CUSTOMER_ID and values match TDR BANs.")
    doc.add_paragraph("SharePoint dropdown not shown → Set Azure secrets and ensure admin granted Sites.Read.All.")
    doc.add_paragraph()

    doc.add_heading("10. Support and Repo", level=1)
    doc.add_paragraph("Repository: TDR-mapping. Main files: app.py, tdr_core.py, sharepoint_graph.py. For runbook updates, update RUNBOOK.md or this script and regenerate RUNBOOK.docx to share with the team.")
    doc.add_paragraph()
    add_para(doc, "End of runbook.", bold=True)

    out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "RUNBOOK.docx")
    doc.save(out_path)
    print(f"Created: {out_path}")

if __name__ == "__main__":
    main()
